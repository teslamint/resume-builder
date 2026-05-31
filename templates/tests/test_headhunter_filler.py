"""Smoke tests for headhunter_filler DOCX template manipulation."""
from pathlib import Path

import pytest

from docx import Document

from headhunter_filler import (
    add_run,
    clear_runs,
    insert_paragraph_after,
    set_plain,
    _fill_cover_letter_inject,
    _clear_between_anchors,
)
from resume_builder import calculate_tenure


def _make_doc(texts: list[str]) -> Document:
    doc = Document()
    for t in texts:
        doc.add_paragraph(t)
    return doc


class TestClearBetweenAnchors:
    def test_clears_paragraphs_between_anchors(self):
        doc = _make_doc(["Header", "stale1", "stale2", "Footer"])
        header, _, end_p, _ = _clear_between_anchors(doc, r"Header", r"Footer")
        texts = [p.text for p in doc.paragraphs]
        assert "stale1" not in texts
        assert "stale2" not in texts
        assert "Header" in texts
        assert "Footer" in texts

    def test_returns_none_when_start_not_found(self):
        doc = _make_doc(["Something", "else"])
        result = _clear_between_anchors(doc, r"MISSING", r"else")
        assert result == (None, None, None, None)

    def test_returns_start_with_none_end_when_end_not_found(self):
        doc = _make_doc(["Header", "content", "more"])
        header, h_idx, end_p, end_idx = _clear_between_anchors(doc, r"Header", r"MISSING")
        assert header is not None
        assert end_p is None


class TestFillCoverLetterInject:
    def test_clears_stale_text_with_signature_anchor(self):
        doc = _make_doc([
            "자기소개서",
            "[지원 동기] 여기에 작성",
            "[주요 성과] 여기에 작성",
            "2024년 00월 00일",
            "지원자 : O O O",
        ])
        target_config = {
            "cover_letter": {
                "motivation": "지원 동기 내용입니다.",
                "future_plan": "포부 내용입니다.",
            }
        }
        rules = {}
        _fill_cover_letter_inject(doc, target_config, rules, "맑은 고딕")

        texts = [p.text for p in doc.paragraphs]
        assert "[지원 동기] 여기에 작성" not in texts
        assert "[주요 성과] 여기에 작성" not in texts
        assert "지원 동기 내용입니다." in texts
        assert "포부 내용입니다." in texts
        # Signature anchor preserved
        assert "2024년 00월 00일" in texts

    def test_clears_stale_text_with_privacy_anchor(self):
        doc = _make_doc([
            "자기소개서",
            "template placeholder",
            "개인정보 수집 동의",
        ])
        target_config = {"cover_letter": {"motivation": "동기"}}
        rules = {}
        _fill_cover_letter_inject(doc, target_config, rules, "맑은 고딕")

        texts = [p.text for p in doc.paragraphs]
        assert "template placeholder" not in texts
        assert "동기" in texts
        assert "개인정보 수집 동의" in texts

    def test_clears_to_end_when_no_anchor(self):
        doc = _make_doc([
            "자기소개서",
            "stale placeholder 1",
            "stale placeholder 2",
        ])
        target_config = {"cover_letter": {"motivation": "새 내용"}}
        rules = {}
        _fill_cover_letter_inject(doc, target_config, rules, "맑은 고딕")

        texts = [p.text for p in doc.paragraphs]
        assert "stale placeholder 1" not in texts
        assert "stale placeholder 2" not in texts
        assert "새 내용" in texts

    def test_noop_when_no_cover_letter_content(self):
        doc = _make_doc(["자기소개서", "placeholder", "2024년 01월 01일"])
        _fill_cover_letter_inject(doc, {}, {}, "맑은 고딕")
        texts = [p.text for p in doc.paragraphs]
        assert "placeholder" not in texts


class TestDocxHelpers:
    def test_helpers_are_shared_from_docx_helpers_module(self):
        import headhunter_filler
        from docx_helpers import set_plain as shared_set_plain

        assert headhunter_filler.set_plain is shared_set_plain

    def test_set_plain_replaces_text(self):
        doc = _make_doc(["original"])
        p = doc.paragraphs[0]
        set_plain(p, "replaced", bold=True, font_name="맑은 고딕")
        assert p.text == "replaced"
        assert p.runs[0].bold is True

    def test_insert_paragraph_after(self):
        doc = _make_doc(["first", "third"])
        ref = doc.paragraphs[0]
        insert_paragraph_after(ref, "second", font_name="맑은 고딕")
        texts = [p.text for p in doc.paragraphs]
        assert texts[1] == "second"


class TestSharedResumeDataHelpers:
    def test_calculate_tenure_supports_headhunter_duration_only_format(self):
        result = calculate_tenure(
            "2020.09 ~ 2022.09",
            separator="~",
            include_period=False,
            error_value="",
        )
        assert result == "2년 1개월"

    def test_calculate_tenure_preserves_resume_builder_default_format(self):
        result = calculate_tenure("2020.09 - 2022.09")
        assert result == "2020.09 - 2022.09 (2년 1개월)"

    def test_resume_builder_owns_contact_and_education_parsers(self, tmp_path):
        from resume_builder import _parse_contact, _parse_education

        contact_path = tmp_path / "contact.md"
        contact_path.write_text(
            "# Contact\n\n- Name: 홍길동\n- Email: test@example.com\n- GitHub: github.com/test\n",
            encoding="utf-8",
        )
        education_path = tmp_path / "education.md"
        education_path.write_text(
            "## Seoul Univ\n\n- Period: 2010.03 - 2014.02\n- Major: Computer Science\n- Status: 졸업\n",
            encoding="utf-8",
        )

        assert _parse_contact(contact_path) == {
            "name": "홍길동",
            "email": "test@example.com",
            "github": "github.com/test",
        }
        assert _parse_education(education_path) == [
            {
                "school": "Seoul Univ",
                "period": "2010.03 - 2014.02",
                "major": "Computer Science",
                "status": "졸업",
            }
        ]
