"""Shared DOCX manipulation helpers for resume template fillers."""
import copy
import re

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, RGBColor

DEFAULT_FONT = "맑은 고딕"
DEFAULT_COLOR = RGBColor(0, 0, 0)
SECTION_SIZE = Emu(139700)  # 11pt


def _font_name(mapping: dict) -> str:
    return mapping.get("font", DEFAULT_FONT)


def clear_runs(p):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    for hl in p._element.findall(qn("w:hyperlink")):
        p._element.remove(hl)


def add_run(p, text, bold=False, size=None, font_name=DEFAULT_FONT):
    r = p.add_run(text)
    r.font.name = font_name
    r.font.color.rgb = DEFAULT_COLOR
    r.bold = bold
    if size:
        r.font.size = size
    rpr = r._element.get_or_add_rPr()
    ea = OxmlElement("w:rFonts")
    ea.set(qn("w:eastAsia"), font_name)
    rpr.insert(0, ea)
    return r


def set_label_value(p, label, value, font_name=DEFAULT_FONT):
    clear_runs(p)
    add_run(p, label, bold=True, font_name=font_name)
    add_run(p, "\t", font_name=font_name)
    add_run(p, value, font_name=font_name)


def set_plain(p, text, bold=False, size=None, font_name=DEFAULT_FONT):
    clear_runs(p)
    add_run(p, text, bold=bold, size=size, font_name=font_name)


def delete_paragraph(p):
    elem = p._element
    elem.getparent().remove(elem)


def insert_paragraph_after(ref_p, text, bold=False, font_name=DEFAULT_FONT):
    new_p = OxmlElement("w:p")
    ref_p._element.addnext(new_p)
    from docx.text.paragraph import Paragraph

    para = Paragraph(new_p, ref_p._parent)
    ppr_src = ref_p._element.find(qn("w:pPr"))
    if ppr_src is not None:
        new_ppr = copy.deepcopy(ppr_src)
        new_p.insert(0, new_ppr)
    add_run(para, text, bold=bold, font_name=font_name)
    return para


def insert_list_paragraph_after(ref_p, text, font_name=DEFAULT_FONT):
    para = insert_paragraph_after(ref_p, "", font_name=font_name)
    clear_runs(para)
    add_run(para, text, font_name=font_name)
    return para


def find_paragraph(doc, pattern: str, start_idx: int = 0, end_idx: int | None = None):
    paras = doc.paragraphs
    end = end_idx if end_idx is not None else len(paras)
    rx = re.compile(pattern, re.IGNORECASE)
    for i in range(start_idx, min(end, len(paras))):
        if rx.search(paras[i].text):
            return i, paras[i]
    return None, None


def clear_between(start_p, end_p):
    elem = start_p._element.getnext()
    while elem is not None and elem is not end_p._element:
        next_e = elem.getnext()
        if elem.tag == qn("w:p"):
            elem.getparent().remove(elem)
        elem = next_e


def fill_table_cell(doc, table_index, row, col, text, font_name=DEFAULT_FONT):
    if table_index >= len(doc.tables):
        return
    table = doc.tables[table_index]
    if row >= len(table.rows) or col >= len(table.columns):
        return
    cell = table.cell(row, col)
    for p in cell.paragraphs:
        clear_runs(p)
    if cell.paragraphs:
        first_p = cell.paragraphs[0]
        lines = text.split("\n")
        add_run(first_p, lines[0], font_name=font_name)
        for extra_line in lines[1:]:
            new_p = OxmlElement("w:p")
            first_p._element.addnext(new_p)
            from docx.text.paragraph import Paragraph

            pp = Paragraph(new_p, first_p._parent)
            add_run(pp, extra_line, font_name=font_name)
            first_p = pp
    for extra_p in cell.paragraphs[1:]:
        if not extra_p.text.strip():
            try:
                delete_paragraph(extra_p)
            except Exception:
                pass
