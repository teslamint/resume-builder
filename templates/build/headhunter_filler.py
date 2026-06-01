#!/usr/bin/env python3
"""Generic headhunter resume template filler.

Usage:
    # Analyze a new template → generate mapping YAML
    python3 headhunter_filler.py analyze <template.docx> [-o mapping.yaml]

    # Fill a template using mapping + resume data
    python3 headhunter_filler.py fill --template <template.docx> --mapping <mapping.yaml> \
        --target-config <target.yaml> [-o output.docx]
"""
import argparse
import json
import re
import sys
from datetime import datetime
from pathlib import Path

import yaml
from docx import Document
from docx.oxml.ns import qn

_SCRIPT_DIR = Path(__file__).parent
_PROJECT_ROOT = _SCRIPT_DIR.parent.parent
_PRIVATE_DIR = _PROJECT_ROOT / "private"

sys.path.insert(0, str(_SCRIPT_DIR))
from docx_helpers import (
    DEFAULT_FONT,
    SECTION_SIZE,
    _font_name,
    add_run,
    clear_between,
    clear_runs,
    delete_paragraph,
    fill_table_cell,
    find_paragraph,
    insert_list_paragraph_after,
    insert_paragraph_after,
    set_label_value,
    set_plain,
)
from resume_builder import (
    _load_company,
    _parse_contact,
    _parse_education,
    calculate_tenure,
    extract_company_info_full,
    extract_project_info,
    extract_section,
    filter_content,
)


# ---------------------------------------------------------------------------
# Text-pattern search helpers
# ---------------------------------------------------------------------------

def _find(doc, pattern, start=0):
    rx = re.compile(pattern, re.IGNORECASE)
    for i in range(start, len(doc.paragraphs)):
        if rx.search(doc.paragraphs[i].text):
            return doc.paragraphs[i], i
    return None, None


def _find_all(doc, pattern):
    rx = re.compile(pattern, re.IGNORECASE)
    return [(doc.paragraphs[i], i) for i in range(len(doc.paragraphs)) if rx.search(doc.paragraphs[i].text)]


def _p(rules: dict, key: str, default: str) -> str:
    """Get pattern from fill_rules.patterns, falling back to default."""
    return rules.get("patterns", {}).get(key, default)


def _clear_between_anchors(doc, start_pattern, end_pattern, start_after=0):
    """Find two anchors and delete all paragraphs between them (exclusive).
    Returns (start_para, start_idx, end_para, end_idx) or Nones."""
    start_p, start_idx = _find(doc, start_pattern, start_after)
    if start_p is None or start_idx is None:
        return None, None, None, None
    end_p, end_idx = _find(doc, end_pattern, start_idx + 1)
    if end_p is None:
        return start_p, start_idx, None, None
    clear_between(start_p, end_p)
    return start_p, start_idx, end_p, end_idx


# ---------------------------------------------------------------------------
# Resume data loader
# ---------------------------------------------------------------------------

def _parse_list_file(path: Path) -> list[str]:
    items = []
    if not path.exists():
        return items
    for line in path.read_text(encoding="utf-8").splitlines():
        if line.startswith("- "):
            items.append(line[2:].strip())
    return items


def load_resume_data(exclude_companies: list[str] | None = None) -> dict:
    profile_dir = _PRIVATE_DIR / "profile"

    personal_path = profile_dir / "personal.yaml"
    personal = {}
    if personal_path.exists():
        personal = yaml.safe_load(personal_path.read_text(encoding="utf-8")) or {}

    contact = _parse_contact(profile_dir / "contact.md")

    education = _parse_education(profile_dir / "education.md")
    for extra in personal.get("additional_education", []):
        education.append(extra)

    skills = _parse_list_file(profile_dir / "skills-job.md")
    competencies = _parse_list_file(profile_dir / "core-competencies.md")
    languages = _parse_list_file(profile_dir / "languages.md")

    exclude = set(exclude_companies or [])
    personal_companies = personal.get("companies", {})

    companies_dir = _PRIVATE_DIR / "companies"
    companies = []
    if companies_dir.is_dir():
        for d in sorted(companies_dir.iterdir()):
            if not d.is_dir() or d.name == "CLAUDE.md" or d.name in exclude:
                continue
            comp = _load_company(d, personal_companies)
            if comp:
                companies.append(comp)

    def _sort_key(c):
        try:
            p = c["period"].split(" ~ ")[0].replace(".", "")
            return -int(p)
        except (ValueError, IndexError):
            return 0

    companies.sort(key=_sort_key)

    total_months = 0
    for c in companies:
        if c["employment"] == "인턴":
            continue
        try:
            parts = c["period"].split(" ~ ")
            sp = parts[0].strip().replace(".", "-").split("-")
            ep = parts[1].strip().replace(".", "-").split("-")
            sy, sm = int(sp[0]), int(sp[1])
            if parts[1].strip() in ("현재", "재직중"):
                now = datetime.now()
                ey, em = now.year, now.month
            else:
                ey, em = int(ep[0]), int(ep[1])
            total_months += (ey - sy) * 12 + (em - sm) + 1
        except (ValueError, IndexError):
            pass

    ty, tm = divmod(total_months, 12)
    total_experience = f"{ty}년 {tm}개월" if tm else f"{ty}년"

    return {
        "personal": {
            "name": contact["name"],
            "email": contact["email"],
            "github": contact["github"],
            "birth_date": personal.get("birth_date", ""),
            "address": personal.get("address", ""),
            "phone": personal.get("phone", ""),
            "military": personal.get("military", ""),
        },
        "education": education,
        "skills": skills,
        "competencies": competencies,
        "languages": languages,
        "companies": companies,
        "total_experience": total_experience,
        "career_year": f"{ty + (1 if tm else 0)}년차",
    }


# ---------------------------------------------------------------------------
# Template analyzer
# ---------------------------------------------------------------------------

SECTION_PATTERNS = {
    "position_header": r"지원.*회사|Position",
    "personal_name": r"성\s*명",
    "personal_birth": r"생\s*년\s*월\s*일",
    "personal_address": r"주\s*소",
    "personal_military": r"병\s*역\s*사\s*항",
    "personal_salary": r"현재\s*직급.*연봉",
    "personal_desired": r"희망\s*직급.*연봉",
    "personal_available": r"입사\s*가능",
    "education_header": r"Education|학\s*력\s*사\s*항",
    "career_header": r"Work\s*Experience|경\s*력\s*사\s*항",
    "core_header": r"핵심\s*경험|Core\s*Experience|주요\s*역량",
    "others_language": r"어학\s*사항",
    "others_cert": r"자격증",
    "others_computer": r"컴퓨터\s*능력",
    "company_intro": r"□?\s*회사\s*소개",
    "company_duties": r"\[?\s*상세\s*담당\s*업무\s*\]?|담당\s*업무",
    "company_achievements": r"\[?\s*주요\s*성과\s*\]?",
    "company_resign": r"\[?\s*퇴직\s*(또는)?\s*이직\s*사유\s*\]?",
    "cover_letter_header": r"자기\s*소개서",
    "cover_motivation": r"\[?\s*지원\s*동기\s*\]?",
    "cover_achievements": r"\[?\s*주요\s*성과\s*(및)?\s*실적\s*\]?",
    "cover_future": r"\[?\s*입사\s*후?\s*포부\s*\]?",
    "signature_date": r"\d{4}년\s*\d{0,2}월?\s*\d{0,2}일?",
    "signature_name": r"지원자\s*[:：]",
}


def analyze_template(docx_path: str) -> dict:
    doc = Document(docx_path)
    detected: dict[str, list[dict]] = {}
    guide_text_candidates: list[dict] = []

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue

        for section_key, pattern in SECTION_PATTERNS.items():
            if re.search(pattern, text, re.IGNORECASE):
                detected.setdefault(section_key, []).append({"idx": i, "text": text})

        if any(
            kw in text
            for kw in [
                "★",
                "ex)",
                "기본적으로",
                "본인의 경력",
                "기억에 남는",
                "경력기술서는",
                "분량",
            ]
        ):
            guide_text_candidates.append({"idx": i, "text": text[:60]})

    company_slot_count = len(detected.get("company_intro", []))

    mapping = {
        "name": Path(docx_path).stem,
        "font": DEFAULT_FONT,
        "detected_sections": {},
        "guide_text_to_delete": [g["text"] for g in guide_text_candidates],
        "template_company_slots": company_slot_count,
        "insert_extra_companies_before": "자기소개서",
    }

    for key, hits in detected.items():
        mapping["detected_sections"][key] = [
            {"paragraph_index": h["idx"], "text": h["text"]} for h in hits
        ]

    return mapping


# ---------------------------------------------------------------------------
# Template filler — inject-based (fill_rules) and legacy (marker-based)
# ---------------------------------------------------------------------------

def fill_template(
    template_path: str,
    mapping: dict,
    target_config: dict,
    resume: dict,
    output_path: str,
):
    doc = Document(template_path)
    font = _font_name(mapping)
    rules = mapping.get("fill_rules", {})

    companies = resume["companies"]
    exclude_employment = target_config.get("exclude_employment_types", ["인턴"])
    display_companies = [c for c in companies if c["employment"] not in exclude_employment]

    if rules:
        _fill_inject(doc, mapping, target_config, resume, display_companies, rules, font)
    else:
        _fill_legacy(doc, mapping, target_config, resume, display_companies, font)

    _delete_guide_text(doc, mapping)
    doc.save(output_path)
    print(f"Saved: {output_path}")


# ---------------------------------------------------------------------------
# Inject-based fill (fill_rules present)
# ---------------------------------------------------------------------------

def _fill_inject(doc, mapping, target_config, resume, display_companies, rules, font):
    _fill_apply_date(doc, rules, font)
    _fill_personal_dispatch(doc, resume, target_config, rules, font)
    _fill_table_sections(doc, resume, target_config, display_companies, rules, font)
    _fill_education_inject(doc, resume, rules, font)
    _fill_career_summary_inject(doc, resume, display_companies, target_config, rules, font)
    _fill_company_details_inject(doc, display_companies, rules, font)
    _fill_core_experiences_inject(doc, display_companies, target_config, rules, font)
    _fill_others_inject(doc, resume, target_config, rules, font)
    _fill_military_inject(doc, resume, rules, font)
    _fill_salary_inject(doc, target_config, resume, rules, font)
    _fill_cover_letter_inject(doc, target_config, rules, font)
    _fill_signature_inject(doc, resume, rules, font)
    _fill_extra_sections(doc, target_config, resume, rules, font)


def _fill_extra_sections(doc, target_config, resume, rules, font):
    for section in rules.get("extra_sections", []):
        anchor = section.get("anchor", "")
        end_anchor = section.get("end_anchor", "")
        config_key = section.get("config_key", "")
        if not anchor or not config_key:
            continue

        header, h_idx, end_p, _ = _clear_between_anchors(doc, anchor, end_anchor)
        if header is None:
            continue

        header_text = section.get("header_text", "")
        if header_text:
            set_plain(header, header_text, bold=True, size=SECTION_SIZE, font_name=font)

        content = target_config.get(config_key)
        if content is None:
            continue

        ref = header
        if isinstance(content, list):
            for item in content:
                ref = insert_paragraph_after(ref, item, font_name=font)
        elif isinstance(content, dict):
            for label, value in content.items():
                ref = insert_paragraph_after(ref, f"{label}\t{value}", font_name=font)
        elif isinstance(content, str):
            for line in content.strip().split("\n"):
                if line.strip():
                    ref = insert_paragraph_after(ref, line.strip(), font_name=font)


def _fill_apply_date(doc, rules, font):
    pattern = _p(rules, "apply_date", r"지원일자")
    if not pattern:
        return
    p, _ = _find(doc, pattern)
    if p:
        today = datetime.now()
        set_plain(p, f"지원일자 : {today.year}년 {today.month}월 {today.day}일", font_name=font)


def _fill_personal_dispatch(doc, resume, target_config, rules, font):
    personal_rules = rules.get("personal", {})
    if personal_rules.get("type") == "table":
        _fill_personal_table(doc, resume, target_config, personal_rules, font)
    else:
        _fill_personal_paragraph(doc, resume, target_config, rules, font)


def _fill_personal_table(doc, resume, target_config, personal_rules, font):
    ti = personal_rules.get("table_index", 0)
    cells = personal_rules.get("cells", {})
    personal = resume["personal"]

    field_values = {
        "position": f"{target_config.get('target_company', '')} / {target_config.get('target_position', '')}",
        "name": personal["name"],
        "birth_date": personal["birth_date"],
        "phone": personal["phone"],
        "email": personal["email"],
        "address": personal["address"],
    }

    current_title = target_config.get("current_title", "")
    current_salary = target_config.get("current_salary", "")
    desired_salary = target_config.get("desired_salary", "면접 후 협의")
    field_values["salary"] = f"최종연봉: {current_salary}\n희망연봉: {desired_salary}"

    for field_name, cell_ref in cells.items():
        if not isinstance(cell_ref, list) or len(cell_ref) != 2:
            continue
        row, col = cell_ref
        value = field_values.get(field_name, "")
        if value:
            fill_table_cell(doc, ti, row, col, value, font_name=font)


def _fill_personal_paragraph(doc, resume, target_config, rules, font):
    personal = resume["personal"]
    personal_rules = rules.get("personal", {})

    custom_fields = personal_rules.get("fields")
    if custom_fields:
        _fill_personal_custom_fields(doc, resume, target_config, custom_fields, font)
        return

    position_pat = _p(rules, "position", r"지원.*회사.*Position")
    p, _ = _find(doc, position_pat)
    if p:
        tc = target_config.get("target_company", "")
        tp = target_config.get("target_position", "")
        set_plain(p, f"지원 회사 / Position : {tc} / {tp}", bold=True, size=SECTION_SIZE, font_name=font)

    field_map = [
        (_p(rules, "name", r"성\s*명"), "성       명", personal["name"]),
        (_p(rules, "birth", r"생\s*년\s*월"), "생 년 월 일", personal["birth_date"]),
        (_p(rules, "address", r"주\s{2,}소"), "주       소", personal["address"]),
        (_p(rules, "military", r"병\s*역"), "병 역 사 항", personal["military"]),
    ]
    for pattern, label, value in field_map:
        p, _ = _find(doc, pattern)
        if p:
            set_label_value(p, label, value, font_name=font)

    current_title = target_config.get("current_title", resume["companies"][0]["role"] if resume["companies"] else "")
    current_salary = target_config.get("current_salary", "")
    p, _ = _find(doc, _p(rules, "current_salary", r"현재\s*직급"))
    if p:
        set_label_value(p, "현재직급/연봉", f"{current_title} / {current_salary}", font_name=font)

    p, _ = _find(doc, _p(rules, "desired_salary", r"희망\s*직급"))
    if p:
        set_label_value(p, "희망직급/연봉", target_config.get("desired_salary", "면접 후 협의"), font_name=font)

    p, _ = _find(doc, _p(rules, "available", r"입사\s*가능"))
    if p:
        set_label_value(p, "입사가능 시기", target_config.get("available_date", "즉시 가능"), font_name=font)


def _fill_personal_custom_fields(doc, resume, target_config, fields, font):
    personal = resume["personal"]
    sources = {
        "name": personal["name"],
        "birth_date": personal["birth_date"],
        "address": personal["address"],
        "phone": personal["phone"],
        "email": personal["email"],
        "military": personal["military"],
        "github": personal.get("github", ""),
        "career": resume.get("total_experience", ""),
        "career_year": resume.get("career_year", ""),
        "position": target_config.get("target_position", ""),
        "target_company": target_config.get("target_company", ""),
        "current_salary": target_config.get("current_salary", ""),
        "desired_salary": target_config.get("desired_salary", ""),
        "available": target_config.get("available_date", ""),
        "name_position": f"{personal['name']} {target_config.get('target_position', '')}",
    }

    for field_def in fields:
        pattern = field_def.get("pattern", "")
        source = field_def.get("source", "")
        label = field_def.get("label", "")
        fmt = field_def.get("format", "")
        if not pattern:
            continue

        p, _ = _find(doc, pattern)
        if not p:
            continue

        if fmt:
            value = fmt.format(**sources)
        else:
            value = sources.get(source, "")

        if label:
            set_label_value(p, label, value, font_name=font)
        else:
            set_plain(p, value, bold=field_def.get("bold", False), font_name=font)


def _fill_table_sections(doc, resume, target_config, display_companies, rules, font):
    for entry in rules.get("table_cells", []):
        ti = entry.get("table", 0)
        row = entry.get("row")
        col = entry.get("col")
        key = entry.get("key", "")
        if row is None or col is None:
            continue

        if key == "core_experiences":
            items = target_config.get("core_experiences", [])
            if not items:
                items = _auto_core_experiences(display_companies)
            text = "\n".join(f"• {item}" for item in items)
            fill_table_cell(doc, ti, row, col, text, font_name=font)

        elif key == "achievements_summary":
            items = target_config.get("achievements_summary", [])
            if not items:
                all_achievements = []
                for c in display_companies[:3]:
                    all_achievements.extend(c.get("achievement_bullets", [])[:2])
                items = all_achievements[:5]
            text = "\n".join(f"• {item}" for item in items)
            fill_table_cell(doc, ti, row, col, text, font_name=font)


def _fill_education_inject(doc, resume, rules, font):
    edu_pat = _p(rules, "education_header", r"Education|학\s*력\s*사\s*항")
    end_pat = _p(rules, "education_end", r"경\s*력\s*사\s*항")

    header, h_idx, end_p, _ = _clear_between_anchors(doc, edu_pat, end_pat)
    if header is None:
        return

    header_fmt = rules.get("education_header_format", "")
    edu = resume["education"]
    if header_fmt and edu:
        main_edu = edu[0]
        fmt_text = header_fmt.format(
            school=main_edu.get("school", ""),
            status=main_edu.get("status", ""),
        )
        set_plain(header, fmt_text, bold=True, size=SECTION_SIZE, font_name=font)

    ref = header
    for e in edu:
        period = e["period"].replace(" - ", " ~ ")
        line = f"{period}   {e['school']} {e.get('major', '')} / {e.get('status', '')}"
        ref = insert_paragraph_after(ref, line, font_name=font)


def _fill_career_summary_inject(doc, resume, display_companies, target_config, rules, font):
    career_pat = _p(rules, "career_header", r"Work\s*Experience|경\s*력\s*사\s*항")
    end_pat = _p(rules, "career_end", r"상세\s*경력|세부\s*경력|핵심.*역량|Core")

    header, h_idx, end_p, _ = _clear_between_anchors(doc, career_pat, end_pat)
    if header is None:
        return

    te = resume["total_experience"]
    re_exp = target_config.get("related_experience", te)
    header_fmt = rules.get("career_header_format", "경력 사항 (총 경력 {total_exp}, 관련경력 {related_exp})")
    set_plain(header, header_fmt.format(total_exp=te, related_exp=re_exp), bold=True, size=SECTION_SIZE, font_name=font)

    ref = header
    for comp in display_companies:
        row_text = f"{comp['period']}\t({comp['tenure']})\t{comp['name']}\t{comp['department']}\t\t{comp['role']}"
        ref = insert_paragraph_after(ref, row_text, font_name=font)


def _fill_company_details_inject(doc, display_companies, rules, font):
    company_pat = _p(rules, "company_header", r"상세\s*경력\s*사항|세부\s*경력\s*사항")
    end_pat = _p(rules, "company_end", r"핵심역량.*강점|외국어.*자격|자기\s*소개서")

    header, h_idx, end_p, _ = _clear_between_anchors(doc, company_pat, end_pat)
    if header is None:
        return

    set_plain(header, "상세 경력 사항", bold=True, size=SECTION_SIZE, font_name=font)

    ref = header
    for ci, comp in enumerate(display_companies):
        if ci > 0:
            ref = insert_paragraph_after(ref, "", font_name=font)

        header_text = f"■ {comp['name']}  {comp['department']} / {comp['role']}    [{comp['period']} ({comp['tenure']})]"
        ref = insert_paragraph_after(ref, header_text, bold=True, font_name=font)

        intro_lines = list(comp.get("intro", []))
        if comp.get("tech_stack_summary"):
            intro_lines.append(f"기술 스택: {comp['tech_stack_summary']}")
        for line in intro_lines:
            ref = insert_paragraph_after(ref, f"  {line}", font_name=font)

        ref = insert_paragraph_after(ref, "", font_name=font)
        ref = insert_paragraph_after(ref, "○ 담당업무", bold=True, font_name=font)
        if comp.get("duties_summary"):
            ref = insert_paragraph_after(ref, f"  {comp['duties_summary']}", font_name=font)
        for bullet in comp.get("duties_bullets", []):
            ref = insert_paragraph_after(ref, f"  • {bullet}", font_name=font)

        ref = insert_paragraph_after(ref, "", font_name=font)
        ref = insert_paragraph_after(ref, "○ 주요성과", bold=True, font_name=font)
        for bullet in comp.get("achievement_bullets", []):
            ref = insert_paragraph_after(ref, f"  • {bullet}", font_name=font)

        ref = insert_paragraph_after(ref, "", font_name=font)
        ref = insert_paragraph_after(ref, f"퇴사 사유 : {comp.get('resign_reason', '')}", font_name=font)


def _fill_core_experiences_inject(doc, display_companies, target_config, rules, font):
    core_pat = _p(rules, "core_header", r"핵심역량.*강점|Core\s*Experience|핵심.*경험")
    end_pat = _p(rules, "core_end", r"외국어.*자격|어학|자기\s*소개서")

    header, h_idx, end_p, _ = _clear_between_anchors(doc, core_pat, end_pat)
    if header is None:
        return

    core_items = target_config.get("core_experiences", [])
    if not core_items:
        core_items = _auto_core_experiences(display_companies)

    ref = header
    for item in core_items:
        ref = insert_paragraph_after(ref, f"• {item}", font_name=font)


def _fill_others_inject(doc, resume, target_config, rules, font):
    others_pat = _p(rules, "others_header", r"외국어.*자격.*교육|어학\s*사항")
    end_pat = _p(rules, "others_end", r"병역사항|병\s*역\s*사\s*항")

    header, h_idx, end_p, _ = _clear_between_anchors(doc, others_pat, end_pat)
    if header is None:
        return

    ref = header

    ref = insert_paragraph_after(ref, "", font_name=font)
    langs = ", ".join(resume["languages"]) if resume["languages"] else "해당 없음"
    ref = insert_paragraph_after(ref, f"외국어: {langs}", font_name=font)

    ref = insert_paragraph_after(ref, "", font_name=font)
    ref = insert_paragraph_after(ref, f"자격사항: {target_config.get('certificates', '해당 없음')}", font_name=font)

    ref = insert_paragraph_after(ref, "", font_name=font)
    skill_summary = target_config.get(
        "computer_skills",
        ", ".join(resume.get("skills", [])[:6]) if resume.get("skills") else "",
    )
    ref = insert_paragraph_after(ref, f"컴퓨터: {skill_summary}", font_name=font)

    if resume["personal"].get("github"):
        ref = insert_paragraph_after(ref, "", font_name=font)
        ref = insert_paragraph_after(ref, f"GitHub: {resume['personal']['github']}", font_name=font)


def _fill_military_inject(doc, resume, rules, font):
    mil_pat = _p(rules, "military", r"병역사항|병\s*역\s*사\s*항")
    if not mil_pat:
        return
    end_pat = _p(rules, "military_end", r"연봉.*입사|현재\s*직급.*연봉")
    if not end_pat:
        return

    header, h_idx, end_p, _ = _clear_between_anchors(doc, mil_pat, end_pat)
    if header is None:
        return

    military = resume["personal"].get("military", "")
    if military:
        ref = insert_paragraph_after(header, military, font_name=font)


def _fill_salary_inject(doc, target_config, resume, rules, font):
    sal_pat = _p(rules, "salary_header", r"연봉.*입사가능|현재\s*직급.*연봉")
    end_pat = _p(rules, "salary_end", r"자기\s*소개서")

    header, h_idx, end_p, _ = _clear_between_anchors(doc, sal_pat, end_pat)
    if header is None:
        return

    set_plain(header, "연봉 및 입사가능시기", bold=True, size=SECTION_SIZE, font_name=font)

    current_title = target_config.get("current_title", "")
    current_salary = target_config.get("current_salary", "")
    desired_salary = target_config.get("desired_salary", "면접 후 협의")
    available = target_config.get("available_date", "즉시 가능")

    ref = header
    ref = insert_paragraph_after(ref, f"현재 연봉 : {current_salary} ({current_title})", font_name=font)
    ref = insert_paragraph_after(ref, f"희망 연봉 : {desired_salary}", font_name=font)
    ref = insert_paragraph_after(ref, f"입사가능시기 : {available}", font_name=font)


def _fill_cover_letter_inject(doc, target_config, rules, font):
    cover_pat = _p(rules, "cover_letter", r"자기\s*소개서")
    end_pat = _p(rules, "cover_letter_end",
                 r"개인정보.*동의|동의함|20\d{2}년\s*\d{1,2}월\s*\d{1,2}일|지원자\s*[:：]")

    header, h_idx, end_p, _ = _clear_between_anchors(doc, cover_pat, end_pat)
    if header is None:
        return

    if end_p is None:
        elem = header._element.getnext()
        while elem is not None:
            next_e = elem.getnext()
            if elem.tag == qn("w:p"):
                elem.getparent().remove(elem)
            elem = next_e

    set_plain(header, "자기소개서", bold=True, size=SECTION_SIZE, font_name=font)

    cover = target_config.get("cover_letter", {})
    if not cover:
        return

    sections = rules.get("cover_letter_sections", [
        {"pattern": "지원\\s*동기", "key": "motivation"},
        {"pattern": "주요\\s*성과", "key": "achievements"},
        {"pattern": "입사\\s*후?\\s*포부", "key": "future_plan"},
    ])

    ref = header
    for section_def in sections:
        config_key = section_def.get("key", "")
        section_title = section_def.get("title", config_key)
        content = cover.get(config_key, "")
        if not content:
            continue

        ref = insert_paragraph_after(ref, "", font_name=font)
        ref = insert_paragraph_after(ref, section_title, bold=True, font_name=font)

        paragraphs = [line.strip() for line in content.strip().split("\n") if line.strip()]
        for para_text in paragraphs:
            ref = insert_paragraph_after(ref, para_text, font_name=font)


def _fill_signature_inject(doc, resume, rules, font):
    date_pat = _p(rules, "signature_date", r"20\d{2}년\s*0{1,2}월\s*0{1,2}일")
    if date_pat:
        for p in doc.paragraphs:
            if re.search(date_pat, p.text):
                today = datetime.now()
                set_plain(p, f"{today.year}년 {today.month:02d}월 {today.day:02d}일", font_name=font)
                break

    name_pat = _p(rules, "signature_name", r"지원자\s*[:：].*O\s*O\s*O|성\s*명\s*[:：].*[(（]?인")
    if name_pat:
        name = resume["personal"]["name"]
        for p in doc.paragraphs:
            if re.search(name_pat, p.text.strip()):
                spaced = " ".join(name) if len(name) <= 4 else name
                if "지원자" in p.text:
                    set_plain(p, f"지원자 : {spaced}", font_name=font)
                else:
                    set_plain(p, f"{spaced}", font_name=font)
                break


# ---------------------------------------------------------------------------
# Legacy fill (no fill_rules — marker-based, backward compatible)
# ---------------------------------------------------------------------------

def _fill_legacy(doc, mapping, target_config, resume, display_companies, font):
    _fill_position_legacy(doc, target_config, font)
    _fill_personal_legacy(doc, resume, target_config, font)
    _fill_education_legacy(doc, resume, font)
    _fill_career_summary_legacy(doc, resume, display_companies, target_config, font)
    _fill_core_experiences_legacy(doc, display_companies, target_config, font)
    _fill_others_legacy(doc, resume, target_config, font)
    _fill_company_details_legacy(doc, mapping, display_companies, font)
    _fill_cover_letter_legacy(doc, target_config, font)
    _fill_signature_legacy(doc, resume, font)


def _fill_position_legacy(doc, target_config, font):
    p, _ = _find(doc, r"지원.*회사.*Position")
    if p:
        tc = target_config.get("target_company", "")
        tp = target_config.get("target_position", "")
        set_plain(p, f"지원 회사 / Position : {tc} / {tp}", bold=True, size=SECTION_SIZE, font_name=font)


def _fill_personal_legacy(doc, resume, target_config, font):
    personal = resume["personal"]
    field_map = [
        (r"성\s*명", "성       명", personal["name"]),
        (r"생\s*년\s*월", "생 년 월 일", personal["birth_date"]),
        (r"주\s{2,}소", "주       소", personal["address"]),
        (r"병\s*역", "병 역 사 항", personal["military"]),
    ]
    for pattern, label, value in field_map:
        p, _ = _find(doc, pattern)
        if p:
            set_label_value(p, label, value, font_name=font)

    current_title = target_config.get("current_title", resume["companies"][0]["role"] if resume["companies"] else "")
    current_salary = target_config.get("current_salary", "")
    p, _ = _find(doc, r"현재\s*직급")
    if p:
        set_label_value(p, "현재직급/연봉", f"{current_title} / {current_salary}", font_name=font)

    p, _ = _find(doc, r"희망\s*직급")
    if p:
        set_label_value(p, "희망직급/연봉", target_config.get("desired_salary", "면접 후 협의"), font_name=font)

    p, _ = _find(doc, r"입사\s*가능")
    if p:
        set_label_value(p, "입사가능 시기", target_config.get("available_date", "즉시 가능"), font_name=font)


def _fill_education_legacy(doc, resume, font):
    header, idx = _find(doc, r"Education|학\s*력\s*사\s*항")
    if not header or idx is None:
        return
    edu = resume["education"]
    main_edu = edu[0] if edu else {}
    set_plain(header, f"Education ({main_edu.get('school', '')} {main_edu.get('status', '')})", bold=True, size=SECTION_SIZE, font_name=font)

    row_idx = idx + 1
    edu_written = 0
    while row_idx < len(doc.paragraphs) and edu_written < len(edu):
        p = doc.paragraphs[row_idx]
        text = p.text.strip()
        if not text or any(kw in text for kw in ["석사", "학사", "고등학교", "대학", "0000"]):
            e = edu[edu_written]
            period = e["period"].replace(" - ", " ~ ")
            set_plain(p, f"{period}\t{e['school']}\t\t{e['major']}\t\t{e['status']}", font_name=font)
            edu_written += 1
            row_idx += 1
        else:
            break

    while row_idx < len(doc.paragraphs):
        p = doc.paragraphs[row_idx]
        text = p.text.strip()
        if not text or any(kw in text for kw in ["석사", "학사", "고등학교", "대학", "0000"]):
            delete_paragraph(p)
        else:
            break


def _fill_career_summary_legacy(doc, resume, display_companies, target_config, font):
    header, idx = _find(doc, r"Work\s*Experience|경\s*력\s*사\s*항")
    if not header or idx is None:
        return

    te = resume["total_experience"]
    re_exp = target_config.get("related_experience", te)
    set_plain(header, f"Work Experience (총 경력 {te}, 관련경력 {re_exp})", bold=True, size=SECTION_SIZE, font_name=font)

    next_p_idx = idx + 1
    if next_p_idx < len(doc.paragraphs):
        gt = doc.paragraphs[next_p_idx].text.strip()
        if "최근" in gt or "상단" in gt or not gt:
            delete_paragraph(doc.paragraphs[next_p_idx])

    template_rows = []
    for si in range(idx + 1, min(idx + 20, len(doc.paragraphs))):
        text = doc.paragraphs[si].text.strip()
        if re.search(r"(핵심|Core|주요\s*역량|어학|자격|Others|세부)", text, re.IGNORECASE):
            break
        if text and ("0000" in text or "재직" in text or re.search(r"\d{4}\.\d{2}", text)):
            template_rows.append(doc.paragraphs[si])

    last_ref = template_rows[-1] if template_rows else header
    for ri, comp in enumerate(display_companies):
        row_text = f"{comp['period']}\t({comp['tenure']})\t{comp['name']}\t{comp['department']}\t\t{comp['role']}"
        if ri < len(template_rows):
            set_plain(template_rows[ri], row_text, font_name=font)
            last_ref = template_rows[ri]
        else:
            last_ref = insert_paragraph_after(last_ref, row_text, font_name=font)

    for leftover in template_rows[len(display_companies):]:
        try:
            delete_paragraph(leftover)
        except Exception:
            pass


def _fill_core_experiences_legacy(doc, display_companies, target_config, font):
    header, idx = _find(doc, r"Core\s*Experience|핵심.*역량|핵심.*경험")
    if not header or idx is None:
        return

    core_items = target_config.get("core_experiences", [])
    if not core_items:
        core_items = _auto_core_experiences(display_companies)

    row_idx = idx + 1
    written = 0
    while row_idx < len(doc.paragraphs) and written < len(core_items):
        p = doc.paragraphs[row_idx]
        text = p.text.strip()
        if re.search(r"(어학|자격|컴퓨터|Others|GitHub|세부.*경력)", text, re.IGNORECASE):
            break
        if text:
            set_plain(p, core_items[written], font_name=font)
            written += 1
        row_idx += 1

    if written < len(core_items):
        ref = doc.paragraphs[row_idx - 1] if row_idx > idx + 1 else header
        for item in core_items[written:]:
            ref = insert_paragraph_after(ref, item, font_name=font)


def _fill_others_legacy(doc, resume, target_config, font):
    p, _ = _find(doc, r"어학\s*사항")
    if p:
        langs = ", ".join(resume["languages"]) if resume["languages"] else "해당 없음"
        set_plain(p, f"어학사항\t{langs}", font_name=font)

    p, _ = _find(doc, r"자격증")
    if p:
        set_plain(p, f"자격증\t\t{target_config.get('certificates', '해당 없음')}", font_name=font)

    p, _ = _find(doc, r"컴퓨터\s*능력")
    if p:
        skill_summary = target_config.get(
            "computer_skills",
            ", ".join(resume.get("skills", [])[:6]) if resume.get("skills") else "",
        )
        set_plain(p, f"컴퓨터 능력\t{skill_summary}", font_name=font)

    for para in doc.paragraphs:
        text = para.text.strip()
        if re.search(r"GitHub|깃허브", text, re.IGNORECASE) and "github.com" not in text:
            set_plain(para, f"GitHub\t\t{resume['personal']['github']}", font_name=font)
            break

    for para in doc.paragraphs:
        text = para.text.strip()
        if re.search(r"이메일|e-?mail", text, re.IGNORECASE) and "@" not in text:
            set_plain(para, f"이메일\t\t{resume['personal']['email']}", font_name=font)
            break


def _fill_company_details_legacy(doc, mapping, companies, font):
    slots = mapping.get("template_company_slots", 2)
    insert_before = mapping.get("insert_extra_companies_before", "자기소개서")

    detail_header, detail_idx = _find(doc, r"세부\s*경력\s*사항")
    if not detail_header or detail_idx is None:
        return

    company_headers = []
    for si in range(detail_idx + 1, len(doc.paragraphs)):
        text = doc.paragraphs[si].text.strip()
        if re.search(r"자기\s*소개서", text):
            break
        if re.search(r"0000\.00|재직|\d{4}\.\d{2}\s*~\s*\d{4}\.\d{2}", text):
            company_headers.append(doc.paragraphs[si])

    intros = _find_all(doc, r"□\s*회사\s*소개")
    duties_hdrs = _find_all(doc, r"\[\s*상세\s*담당\s*업무\s*\]")
    achiev_hdrs = _find_all(doc, r"\[\s*주요\s*성과\s*\]")
    resign_hdrs = _find_all(doc, r"\[\s*퇴직\s*(또는)?\s*이직\s*사유\s*\]")

    for ci, comp in enumerate(companies[:slots]):
        if ci < len(company_headers):
            header_text = f"{comp['period']} ({comp['tenure']})\t{comp['name']}\t\t{comp['department']}/{comp['role']}"
            set_plain(company_headers[ci], header_text, bold=True, font_name=font)

        if ci < len(intros):
            intro_p = intros[ci][0]
            intro_bullets = list(comp.get("intro", []))
            if comp.get("tech_stack_summary"):
                intro_bullets.append(f"기술 스택: {comp['tech_stack_summary']}")

            elem = intro_p._element.getnext()
            filled = 0
            while elem is not None and filled < len(intro_bullets):
                if elem.tag == qn("w:p"):
                    from docx.text.paragraph import Paragraph
                    pp = Paragraph(elem, intro_p._parent)
                    if re.search(r"(상세.*담당|\[상세)", pp.text, re.IGNORECASE):
                        break
                    set_plain(pp, intro_bullets[filled], font_name=font)
                    filled += 1
                elem = elem.getnext()

        if ci < len(duties_hdrs) and ci < len(achiev_hdrs):
            duties_p = duties_hdrs[ci][0]
            achiev_p = achiev_hdrs[ci][0]
            clear_between(duties_p, achiev_p)

            ref = insert_paragraph_after(duties_p, comp.get("duties_summary", ""), font_name=font)
            for bullet in comp.get("duties_bullets", []):
                ref = insert_list_paragraph_after(ref, bullet, font_name=font)

        if ci < len(achiev_hdrs) and ci < len(resign_hdrs):
            achiev_p = achiev_hdrs[ci][0]
            resign_p = resign_hdrs[ci][0]
            clear_between(achiev_p, resign_p)

            ref = achiev_p
            for bullet in comp.get("achievement_bullets", []):
                ref = insert_paragraph_after(ref, bullet, font_name=font)

        if ci < len(resign_hdrs):
            set_plain(resign_hdrs[ci][0], f"[퇴직 또는 이직사유] {comp.get('resign_reason', '')}", bold=True, font_name=font)

    used = min(len(companies), slots)
    _delete_unused_company_slots(company_headers, intros, duties_hdrs, achiev_hdrs, resign_hdrs, used)

    if len(companies) > slots:
        _insert_extra_companies(doc, companies[slots:], insert_before, font)


def _delete_unused_company_slots(headers, intros, duties, achievs, resigns, used):
    def _elem(entry):
        return entry[0]._element if isinstance(entry, tuple) else entry._element

    max_slots = max(len(headers), len(intros), len(duties), len(achievs), len(resigns))
    for si in range(used, max_slots):
        start_elem = _elem(headers[si]) if si < len(headers) else (_elem(intros[si]) if si < len(intros) else None)
        end_elem = _elem(resigns[si]) if si < len(resigns) else None
        if start_elem is None:
            continue
        parent = start_elem.getparent()
        if parent is None:
            continue
        if end_elem is None:
            parent.remove(start_elem)
            continue
        elem = start_elem
        while elem is not None:
            nxt = elem.getnext()
            is_end = elem is end_elem
            if elem.tag == qn("w:p"):
                parent.remove(elem)
            if is_end:
                break
            elem = nxt


def _insert_extra_companies(doc, extra_companies, insert_before_text, font):
    cover_p = None
    for p in doc.paragraphs:
        if insert_before_text in p.text:
            cover_p = p
            break
    if not cover_p:
        return

    prev_elem = cover_p._element.getprevious()
    while prev_elem is not None and prev_elem.tag != qn("w:p"):
        prev_elem = prev_elem.getprevious()
    if prev_elem is None:
        return

    from docx.text.paragraph import Paragraph
    insert_point = Paragraph(prev_elem, cover_p._parent)

    for comp in reversed(extra_companies):
        insert_paragraph_after(insert_point, "", font_name=font)
        insert_paragraph_after(insert_point, f"[퇴직 또는 이직사유] {comp.get('resign_reason', '')}", bold=True, font_name=font)
        insert_paragraph_after(insert_point, "", font_name=font)

        for line in reversed(comp.get("achievement_bullets", [])[:3]):
            insert_paragraph_after(insert_point, line, font_name=font)
        insert_paragraph_after(insert_point, "[주요성과]", bold=True, font_name=font)
        insert_paragraph_after(insert_point, "", font_name=font)

        for bullet in reversed(comp.get("duties_bullets", [])[:4]):
            insert_paragraph_after(insert_point, "• " + bullet, font_name=font)
        if comp.get("duties_summary"):
            insert_paragraph_after(insert_point, comp["duties_summary"], font_name=font)
        insert_paragraph_after(insert_point, "[상세 담당업무]", bold=True, font_name=font)
        insert_paragraph_after(insert_point, "", font_name=font)

        intro_with_tech = list(comp.get("intro", []))
        if comp.get("tech_stack_summary"):
            intro_with_tech.append(f"기술 스택: {comp['tech_stack_summary']}")
        for line in reversed(intro_with_tech):
            insert_paragraph_after(insert_point, "• " + line, font_name=font)
        insert_paragraph_after(insert_point, "□ 회사소개", bold=True, font_name=font)

        header_text = f"{comp['period']} ({comp['tenure']})\t{comp['name']}\t\t{comp['department']}/{comp['role']}"
        insert_paragraph_after(insert_point, header_text, bold=True, font_name=font)
        insert_paragraph_after(insert_point, "", font_name=font)


def _fill_cover_letter_legacy(doc, target_config, font):
    cover = target_config.get("cover_letter", {})
    if not cover:
        return

    cover_fields = [
        (r"\[\s*지원\s*동기\s*\]", "motivation"),
        (r"\[\s*주요\s*성과\s*(및)?\s*실적\s*\]", "achievements"),
        (r"\[\s*입사\s*후?\s*포부\s*\]", "future_plan"),
    ]

    for pattern, config_key in cover_fields:
        if config_key not in cover:
            continue
        p, idx = _find(doc, pattern)
        if not p or idx is None:
            continue

        content = cover[config_key].strip()
        paragraphs = [line.strip() for line in content.split("\n") if line.strip()]
        if not paragraphs:
            continue

        guide_idx = idx + 1
        if guide_idx < len(doc.paragraphs):
            set_plain(doc.paragraphs[guide_idx], paragraphs[0], font_name=font)
            ref = doc.paragraphs[guide_idx]
            for para_text in paragraphs[1:]:
                ref = insert_paragraph_after(ref, para_text, font_name=font)


def _fill_signature_legacy(doc, resume, font):
    for p in doc.paragraphs:
        if re.search(r"20\d{2}년\s*0{1,2}월\s*0{1,2}일", p.text):
            today = datetime.now()
            set_plain(p, f"{today.year}년 {today.month:02d}월 {today.day:02d}일", font_name=font)
            break

    for p in doc.paragraphs:
        if re.search(r"지원자\s*[:：].*O\s*O\s*O", p.text):
            name = resume["personal"]["name"]
            set_plain(p, f"지원자 : {' '.join(name)}", font_name=font)
            break


# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------

def _auto_core_experiences(companies: list[dict]) -> list[str]:
    items = []
    for comp in companies[:3]:
        summary = comp.get("duties_summary", "")
        if summary:
            items.append(f"{comp['name']}: {summary}")
    return items[:5]


def _delete_guide_text(doc, mapping):
    patterns = mapping.get("guide_text_to_delete", [])
    if not patterns:
        return
    for p in list(doc.paragraphs):
        text = p.text.strip()
        for pat in patterns:
            if pat in text:
                try:
                    delete_paragraph(p)
                except Exception:
                    pass
                break


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def cmd_analyze(args):
    mapping = analyze_template(args.template)
    output = args.output or Path(args.template).stem + "_mapping.yaml"
    with open(output, "w", encoding="utf-8") as f:
        yaml.dump(mapping, f, allow_unicode=True, default_flow_style=False, sort_keys=False)
    print(f"Mapping written to: {output}")
    print(f"Detected {len(mapping['detected_sections'])} section types, "
          f"{mapping['template_company_slots']} company slot(s)")
    print("\nReview the YAML and adjust as needed, then use 'fill' command.")


def cmd_fill(args):
    with open(args.mapping, "r", encoding="utf-8") as f:
        mapping = yaml.safe_load(f)

    with open(args.target_config, "r", encoding="utf-8") as f:
        target_config = yaml.safe_load(f)

    exclude = target_config.get("exclude_companies", [])
    resume = load_resume_data(exclude_companies=exclude)

    template = args.template or target_config.get("template_path", "")
    if not template:
        print("Error: --template or template_path in target config required", file=sys.stderr)
        sys.exit(1)

    output = args.output or target_config.get("output_path", "resume-headhunter.docx")

    fill_template(template, mapping, target_config, resume, output)


def cmd_dump_data(args):
    resume = load_resume_data()
    print(json.dumps(resume, ensure_ascii=False, indent=2))


def main():
    parser = argparse.ArgumentParser(description="Generic headhunter resume template filler")
    sub = parser.add_subparsers(dest="command", required=True)

    p_analyze = sub.add_parser("analyze", help="Analyze a DOCX template and generate mapping YAML")
    p_analyze.add_argument("template", help="Path to DOCX template")
    p_analyze.add_argument("-o", "--output", help="Output YAML path")

    p_fill = sub.add_parser("fill", help="Fill a template using mapping + target config")
    p_fill.add_argument("--template", "-t", help="Path to DOCX template")
    p_fill.add_argument("--mapping", "-m", required=True, help="Path to mapping YAML")
    p_fill.add_argument("--target-config", "-c", required=True, help="Path to target config YAML")
    p_fill.add_argument("-o", "--output", help="Output DOCX path")

    p_dump = sub.add_parser("dump-data", help="Dump loaded resume data as JSON")

    args = parser.parse_args()
    if args.command == "analyze":
        cmd_analyze(args)
    elif args.command == "fill":
        cmd_fill(args)
    elif args.command == "dump-data":
        cmd_dump_data(args)


if __name__ == "__main__":
    main()
